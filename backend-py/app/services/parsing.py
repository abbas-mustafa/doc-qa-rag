"""Per-file-type document parsing.

PDF  -> PyMuPDF: text per page, tables (find_tables), embedded figures (vision),
        and OCR (vision) for scanned/image-only pages.
DOCX -> python-docx: structure-aware text (headings preserved), tables, images (vision).
TXT  -> plain read.

CPU-bound extraction runs in a worker thread; Gemini vision calls run concurrently
(bounded) in the async layer.
"""
from __future__ import annotations

import asyncio
import logging
from dataclasses import dataclass, field

import fitz  # PyMuPDF
import docx

from . import vision

logger = logging.getLogger("docqa.parsing")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Tuning
_MIN_IMAGE_DIM = 100          # skip icons / bullets / dividers
_MAX_FIGURES_PER_DOC = 30     # cap vision cost per document
_MAX_OCR_PAGES = 40
_SCANNED_TEXT_THRESHOLD = 40  # chars; below this + has image => treat page as scanned
_VISION_CONCURRENCY = 2  # gentle on Gemini free-tier per-minute limits


@dataclass
class Element:
    text: str
    page_number: int | None
    element_type: str  # text | ocr_text | figure | table


@dataclass
class _RawExtraction:
    text_elements: list[Element] = field(default_factory=list)
    figure_images: list[tuple[int | None, bytes]] = field(default_factory=list)
    ocr_images: list[tuple[int | None, bytes]] = field(default_factory=list)
    page_count: int | None = None


@dataclass
class ParsedDocument:
    elements: list[Element]
    page_count: int | None


# --------------------------------------------------------------------------- PDF


def _pixmap_png(doc: fitz.Document, xref: int) -> bytes | None:
    """Render an image xref to PNG bytes, normalizing colorspace."""
    try:
        pix = fitz.Pixmap(doc, xref)
        if pix.n >= 5:  # CMYK or with alpha -> convert to RGB
            pix = fitz.Pixmap(fitz.csRGB, pix)
        return pix.tobytes("png")
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to render image xref %s: %s", xref, exc)
        return None


def _extract_pdf(path: str) -> _RawExtraction:
    raw = _RawExtraction()
    doc = fitz.open(path)
    raw.page_count = doc.page_count
    seen_image_xrefs: set[int] = set()

    for page_index in range(doc.page_count):
        page = doc[page_index]
        page_no = page_index + 1

        text = page.get_text("text").strip()

        # Tables (each becomes its own element for clean chunking).
        try:
            for table in page.find_tables().tables:
                md = table.to_markdown().strip()
                if md:
                    raw.text_elements.append(Element(md, page_no, "table"))
        except Exception as exc:  # noqa: BLE001
            logger.debug("find_tables failed on page %s: %s", page_no, exc)

        images = page.get_images(full=True)

        # Scanned/image-only page: little extractable text but has imagery -> OCR.
        if len(text) < _SCANNED_TEXT_THRESHOLD and images:
            if len(raw.ocr_images) < _MAX_OCR_PAGES:
                pix = page.get_pixmap(dpi=150)
                raw.ocr_images.append((page_no, pix.tobytes("png")))
        elif text:
            raw.text_elements.append(Element(text, page_no, "text"))

        # Embedded figures worth describing.
        for img in images:
            xref = img[0]
            if xref in seen_image_xrefs:
                continue
            seen_image_xrefs.add(xref)
            if len(raw.figure_images) >= _MAX_FIGURES_PER_DOC:
                break
            width, height = img[2], img[3]
            if width < _MIN_IMAGE_DIM or height < _MIN_IMAGE_DIM:
                continue
            png = _pixmap_png(doc, xref)
            if png:
                raw.figure_images.append((page_no, png))

    doc.close()
    return raw


# -------------------------------------------------------------------------- DOCX


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    header_sep = "| " + " | ".join("---" for _ in table.columns) + " |"
    return "\n".join([rows[0], header_sep, *rows[1:]])


def _extract_docx(path: str) -> _RawExtraction:
    raw = _RawExtraction()
    document = docx.Document(path)

    # Structure-aware text: prefix headings with markdown hashes so the chunker
    # can respect section boundaries.
    lines: list[str] = []
    for para in document.paragraphs:
        content = para.text.strip()
        if not content:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "1"
            lines.append(f"{'#' * min(int(level), 6)} {content}")
        elif style == "title":
            lines.append(f"# {content}")
        else:
            lines.append(content)
    if lines:
        raw.text_elements.append(Element("\n".join(lines), None, "text"))

    for table in document.tables:
        md = _table_to_markdown(table).strip()
        if md:
            raw.text_elements.append(Element(md, None, "table"))

    # Images embedded in the docx package.
    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            try:
                blob = rel.target_part.blob
            except Exception:  # noqa: BLE001
                continue
            if len(raw.figure_images) >= _MAX_FIGURES_PER_DOC:
                break
            raw.figure_images.append((None, blob))

    return raw


# --------------------------------------------------------------------------- TXT


def _extract_txt(path: str) -> _RawExtraction:
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        text = fh.read().strip()
    raw = _RawExtraction(page_count=None)
    if text:
        raw.text_elements.append(Element(text, None, "text"))
    return raw


def _extract_raw(path: str, mime_type: str) -> _RawExtraction:
    if mime_type == "application/pdf":
        return _extract_pdf(path)
    if mime_type == DOCX_MIME:
        return _extract_docx(path)
    if mime_type == "text/plain":
        return _extract_txt(path)
    raise ValueError(f"Unsupported mime type: {mime_type}")


# ----------------------------------------------------------------------- async


async def parse_document(path: str, mime_type: str) -> ParsedDocument:
    raw = await asyncio.to_thread(_extract_raw, path, mime_type)

    semaphore = asyncio.Semaphore(_VISION_CONCURRENCY)

    async def describe(page_no: int | None, image: bytes) -> Element | None:
        async with semaphore:
            desc = await vision.describe_figure(image)
        if not desc:
            return None
        return Element(f"[Figure] {desc}", page_no, "figure")

    async def ocr(page_no: int | None, image: bytes) -> Element | None:
        async with semaphore:
            text = await vision.ocr_page(image)
        if not text:
            return None
        return Element(text, page_no, "ocr_text")

    tasks = [describe(p, img) for p, img in raw.figure_images]
    tasks += [ocr(p, img) for p, img in raw.ocr_images]

    vision_elements: list[Element] = []
    if tasks:
        for result in await asyncio.gather(*tasks, return_exceptions=True):
            if isinstance(result, Element):
                vision_elements.append(result)
            elif isinstance(result, Exception):
                logger.warning("Vision task failed: %s", result)

    elements = raw.text_elements + vision_elements
    return ParsedDocument(elements=elements, page_count=raw.page_count)
